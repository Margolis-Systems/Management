class images:
    def gen_qr(self,data):
        import qrcode
        qr = qrcode.QRCode(
            version=1,
            box_size=10,
            border=5)
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image(fill='black', back_color='white')
        name = 'qrcode001.png'
        img.save('static\\img\\'+name)
        return name

    def gen_pdf417(self,data):
        import pdf417
        _data = pdf417.encode(data)
        img = pdf417.render_image(_data)
        name = 'pdf417002.png'
        img.save('static\\img\\' + name)
        return name

    def create_code_for_qr(self): # not in use
        return "BF2D@Hj @r22071206@i@p1@l4000@n200@e710@d12@g@\nGl4000@w0@\nC68@x".encode('utf-8')

    def create_shape_plot(self, shape, data): # posisiton [x,y]
        import os
        from PIL import Image, ImageDraw, ImageFont
        from pages import shapes
        print(shape)
        print(data)
        print('\n\n')
        positions = shapes[shape]['positions']
        static_dir = os.path.dirname(__file__)+'\\static\\'
        img_dir = static_dir + 'images\\shapes\\' + str(shape) + '.png'
        #img_dir = static_dir + shapes[shape]['img_dir'] #todo:use this after mongo update
        if os.path.exists(img_dir):
            img = Image.open(img_dir)
        else:
            img = Image.open(static_dir + 'images\\shapes\\0.png')
        draw = ImageDraw.Draw(img)
        for index in range(len(data)):
            bbox = draw.textbbox(positions[index], data[str(index + 1)], font = ImageFont.truetype("segoeui.ttf", 14))
            draw.rectangle(bbox, fill="white")
            draw.text(positions[index], data[str(index + 1)], font=ImageFont.truetype("segoeui.ttf", 14), fill="black")

        file_out = static_dir + 'img\\temp_'+str(shape)+'.png'
        img.save(file_out)
        return file_out

class printers:
    def word_print(self,filename=""):
        import os
        word_full_path = "C:\\Program Files\\Microsoft Office\\root\\Office16\\WINWORD.EXE"
        printer_name = "HP8F170F (HP DeskJet 3700 series)"#"Microsoft Print to PDF"
        if not filename:
            filename = "C:\\Users\\baruch.larionov\\Desktop\\test\\testing1.docx"
        os.system('RUNDLL32 PRINTUI.DLL,PrintUIEntry /y /n "'+printer_name+'"&"'+word_full_path+'" /q /n "'+filename+'" /mFilePrintDefault /mFileCloseOrExit')

class reports:
    def fill_word_header(self, dict, src_dir=""):
        from mailmerge import MailMerge
        import os
        from datetime import datetime
        if not src_dir:
            src_dir = "reports\\reports_templates\\orders_template.docx"
        dict['print_date'] = datetime.strftime(datetime.now(),"%d%m%Y_%H%M%S")
        for item in dict:
            dict[item] = str(dict[item])
        document = MailMerge(src_dir)
        document.merge(**dict)
        out_file = os.path.dirname(__file__)+ "\\reports\\reports_temp\\"+dict['id']+"_"+dict['print_date']+".docx" #+ os.path.split(src_dir)[-1]
        document.write(out_file)
        return out_file

    def generate_summary(self,data,weight_list): #todo: finish func
        summary = {}
        total_weight = 0
        for row in data:
            if row['קוטר'] not in summary.keys():
                summary[row['קוטר']] = {'אורך':float(row['אורך']),'משקל':float(row['משקל'])}
            else:
                summary[row['קוטר']]['אורך'] += float(row['אורך'])
                summary[row['קוטר']]['משקל'] += float(row['משקל'])
            total_weight += float(row['משקל'])

        return summary, total_weight

    def generate_order_report(self,template, data, reverse, weight_list, convert_to_pdf=False):
        import docx
        # from docx2pdf import convert
        from docx.enum.table import WD_TABLE_DIRECTION
        import pathlib
        import os
        from datetime import datetime

        template = reports.fill_word_header(self,data['info'],template)
        doc = docx.Document(template)
        table_data = sorted(data['data'], key=lambda x: x['שורה'], reverse=(not reverse))
        table = doc.add_table(len(table_data)*4, 4)
        table.direction = WD_TABLE_DIRECTION.RTL
        table.style = 'Table Grid'
        table.allow_autofit = False
        # Order Table
        key_list = list(table_data[0].keys())
        for i in range(len(table_data)):
            for j in range(5):
                for k in range(5):
                    if j + k == 0:
                        cel = table.rows[i * 4].cells[0]
                        paragraph = cel.paragraphs[0]
                        run = paragraph.add_run()
                        img_file_name = images.create_shape_plot(self,table_data[i]['צורה'],data['shape_data'][i])
                        run.add_picture(img_file_name, width=3500000, height=700000)
                    elif k != 0:
                        index = 3 * j + k - 1
                        if index < len(key_list):
                            table.cell((i * 4 + j), k).text = str(table_data[i][key_list[index]]) + " " + key_list[index]
            table.cell(i * 4, 0).merge(table.cell(i * 4 + 3, 0))
        out_file = os.path.dirname(__file__)+"\\reports\\report_output\\" + os.path.basename(template)
        # Summary
        doc.add_page_break()
        summary, total_weight = reports.generate_summary(self, table_data, weight_list)
        table = doc.add_table(len(summary) + 3, 5)
        table.direction = WD_TABLE_DIRECTION.RTL
        table.style = 'Table Grid'
        table.allow_autofit = False
        key_list = list(summary.keys())
        table.cell(0,0).text = "סיכום משקל ברזל"
        table.cell(0,0).merge(table.cell(0, 4))
        for i in range(1, len(summary)+2):
            if i == 1:
                table.cell(i,0).text = 'משקל תאורטי בק"ג'
                table.cell(i,1).text = 'משקל למטר'
                table.cell(i,2).text = 'סה"כ אורך בס"מ'
                table.cell(i,3).text = 'קוטר ברזל'
                table.cell(i,4).text = 'סוג ברזל'
            else:
                table.cell(i,0).text = str(summary[key_list[i-2]]['משקל'])
                table.cell(i,1).text = str(weight_list[key_list[i-2]])
                table.cell(i,2).text = str(summary[key_list[i-2]]['אורך'])
                table.cell(i,3).text = str(key_list[i-2])
        table.cell(len(summary)+2, 0).text = str(total_weight)
        table.cell(len(summary)+2, 1).merge(table.cell(len(summary)+2, 3))
        table.cell(len(summary)+2, 4).text = "סיכום משקל ברזל"

        doc.save(out_file)
        # if convert_to_pdf:
        #     convert(out_file)
        #     out_file = out_file.replace("docx","pdf")
        return out_file